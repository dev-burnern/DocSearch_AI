from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import openpyxl
from docx import Document
from pypdf import PdfReader


@dataclass(frozen=True)
class TextUnit:
    text: str
    page: int | None = None
    sheet: str | None = None


_CTRL_RE = re.compile(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F]")
_ZW_RE = re.compile(r"[\u200B-\u200D\uFEFF]")


def normalize_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = _CTRL_RE.sub(" ", s)
    s = _ZW_RE.sub("", s)
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def extract_text_units(path: Path) -> list[TextUnit]:
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in (".xlsx", ".xlsm"):
        return _extract_xlsx(path)
    if ext in (".txt", ".md"):
        txt = normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
        return [TextUnit(text=txt)]
    raise ValueError(f"Unsupported file type: {path.name}")


def _extract_pdf(path: Path) -> list[TextUnit]:
    reader = PdfReader(str(path))
    units: list[TextUnit] = []
    for i, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        txt = normalize_text(raw)
        if txt:
            units.append(TextUnit(text=txt, page=i))
    if not units:
        raise RuntimeError(f"PDF text extraction produced empty result (maybe scanned PDF): {path}")
    return units


def _extract_docx(path: Path) -> list[TextUnit]:
    doc = Document(str(path))
    parts: list[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)

    txt = normalize_text("\n\n".join(parts))
    if not txt:
        raise RuntimeError(f"DOCX text extraction produced empty result: {path}")
    return [TextUnit(text=txt)]


def _iter_rows(ws) -> Iterable[str]:
    for row in ws.iter_rows(values_only=True):
        vals = []
        for v in row:
            if v is None:
                continue
            s = str(v).strip()
            if s:
                vals.append(s)
        if vals:
            yield " | ".join(vals)


def _extract_xlsx(path: Path) -> list[TextUnit]:
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    units: list[TextUnit] = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = list(_iter_rows(ws))
        txt = normalize_text("\n".join(rows))
        if txt:
            units.append(TextUnit(text=txt, sheet=name))
    if not units:
        raise RuntimeError(f"XLSX text extraction produced empty result: {path}")
    return units
